"""
Asymmetrica E2E OCR Pipeline — Full Extraction with Optimal Triage
===================================================================
Complete pipeline: Triage → Free Parse → Vector Extract → Mistral OCR → Merge

Phases:
  Phase 0: Discovery + Classification (all files)
  Phase 1: Free extraction (XLSX, DOCX, RTF, MSG, XML — $0)
  Phase 2: Vector PDF extraction (PyMuPDF — $0)
  Phase 3: OCR via Mistral (scanned PDFs + images — ~$0.001/page)
  Phase 4: Per-project merge (all sources → unified project record)

Usage:
  python ocr_pipeline.py \\
    --folder "path/to/offers" \\
    --mistral-key "KEY" \\
    [--ocr]              # Actually run OCR (default: triage only)
    [--ocr-limit N]      # Cap OCR calls for testing (default: unlimited)
    [--json out.json]    # Export full results

Philosophy: "The cheapest OCR call is the one you never make."
"""

import argparse
import base64
import hashlib
import json
import os
import sys
import time
import traceback
from collections import Counter, defaultdict
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import Optional
from concurrent.futures import ThreadPoolExecutor, as_completed

import httpx


# ─── Configuration ────────────────────────────────────────────────

VECTOR_PDF_CHAR_THRESHOLD = 100
MISTRAL_OCR_ENDPOINT = "https://api.mistral.ai/v1/ocr"
MISTRAL_RATE_LIMIT_DELAY = 0.5  # seconds between API calls (be nice)
MAX_FILE_SIZE_FOR_OCR = 50 * 1024 * 1024  # 50 MB cap per file
OCR_TIMEOUT = 60  # seconds per OCR call


# ─── Lane Classification ───────────────────────────────────────────

class Lane(str, Enum):
    FREE = "LANE_FREE"
    VECTOR = "LANE_VECTOR"
    OCR = "LANE_OCR"
    UNPACK = "LANE_UNPACK"
    SKIP = "LANE_SKIP"


EXT_TO_LANE = {
    ".xlsx": Lane.FREE, ".xls": Lane.FREE, ".csv": Lane.FREE,
    ".docx": Lane.FREE, ".doc": Lane.FREE,
    ".rtf": Lane.FREE,
    ".msg": Lane.FREE, ".eml": Lane.FREE,
    ".xml": Lane.FREE,
    ".txt": Lane.FREE, ".json": Lane.FREE,
    ".jpg": Lane.OCR, ".jpeg": Lane.OCR,
    ".png": Lane.OCR, ".bmp": Lane.OCR,
    ".tiff": Lane.OCR, ".tif": Lane.OCR,
    ".zip": Lane.UNPACK, ".rar": Lane.UNPACK, ".7z": Lane.UNPACK,
    ".mp4": Lane.SKIP, ".mp3": Lane.SKIP, ".wav": Lane.SKIP,
    ".html": Lane.SKIP, ".htm": Lane.SKIP,
    ".exe": Lane.SKIP, ".dll": Lane.SKIP,
    ".ini": Lane.SKIP, ".log": Lane.SKIP,
    ".pdf": None,
}


# ─── Data Structures ───────────────────────────────────────────────

@dataclass
class FileResult:
    path: str
    filename: str
    extension: str
    size_bytes: int
    lane: str
    project_folder: str = ""
    subfolder: str = ""
    char_count: int = 0
    page_count: int = 0
    confidence: float = 1.0
    extracted_text: str = ""
    extracted_text_preview: str = ""
    metadata: dict = field(default_factory=dict)
    error: str = ""
    processing_time_ms: float = 0.0
    ocr_model: str = ""  # Which model processed this file


@dataclass
class ProjectRecord:
    """Merged extraction for one offer project (e.g., EH-15-26 ALBA FIT)."""
    project_id: str
    client_name: str = ""
    file_count: int = 0
    total_size_bytes: int = 0
    files: list = field(default_factory=list)
    costing_data: list = field(default_factory=list)
    commercial_offers: list = field(default_factory=list)
    technical_specs: list = field(default_factory=list)
    emails: list = field(default_factory=list)
    rfq_documents: list = field(default_factory=list)
    ocr_results: list = field(default_factory=list)
    subfolders: list = field(default_factory=list)
    confidence_avg: float = 0.0


# ─── File Parsers (Phase 1: FREE) ─────────────────────────────────

def parse_xlsx(filepath: str) -> dict:
    """Extract from modern Excel files (.xlsx)."""
    import openpyxl
    wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
    result = {"sheet_names": wb.sheetnames, "sheets": {}}
    total_chars = 0
    for sheet_name in wb.sheetnames[:5]:
        ws = wb[sheet_name]
        rows = []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i >= 50:
                break
            row_data = [str(c) if c is not None else "" for c in row]
            rows.append(row_data)
            total_chars += sum(len(s) for s in row_data)
        result["sheets"][sheet_name] = {
            "rows_preview": rows,
            "max_row": ws.max_row,
            "max_col": ws.max_column,
        }
    wb.close()
    result["total_chars"] = total_chars
    # Build full text for merging
    all_text = []
    for sname, sdata in result["sheets"].items():
        for row in sdata["rows_preview"]:
            all_text.append(" | ".join(c for c in row if c))
    result["text"] = "\n".join(all_text)
    result["char_count"] = total_chars
    return result


def parse_xls_legacy(filepath: str) -> dict:
    """Extract from legacy Excel files (.xls) using xlrd."""
    import xlrd
    wb = xlrd.open_workbook(filepath)
    result = {"sheet_names": wb.sheet_names(), "sheets": {}}
    total_chars = 0
    for sheet_name in wb.sheet_names()[:5]:
        ws = wb.sheet_by_name(sheet_name)
        rows = []
        for i in range(min(ws.nrows, 50)):
            row_data = [str(ws.cell_value(i, j)) for j in range(ws.ncols)]
            rows.append(row_data)
            total_chars += sum(len(s) for s in row_data)
        result["sheets"][sheet_name] = {
            "rows_preview": rows,
            "max_row": ws.nrows,
            "max_col": ws.ncols,
        }
    all_text = []
    for sname, sdata in result["sheets"].items():
        for row in sdata["rows_preview"]:
            all_text.append(" | ".join(c for c in row if c))
    result["text"] = "\n".join(all_text)
    result["char_count"] = total_chars
    result["total_chars"] = total_chars
    return result


def parse_docx(filepath: str) -> dict:
    """Extract from modern Word documents (.docx)."""
    import docx
    doc = docx.Document(filepath)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)
    tables_data = []
    for table in doc.tables[:5]:
        table_rows = []
        for row in table.rows[:20]:
            table_rows.append([cell.text.strip() for cell in row.cells])
        tables_data.append(table_rows)
    return {
        "text": text,
        "char_count": len(text),
        "paragraph_count": len(paragraphs),
        "table_count": len(doc.tables),
        "tables_preview": tables_data,
    }


def parse_doc_legacy(filepath: str) -> dict:
    """Extract text from legacy .doc files using olefile."""
    import olefile
    if not olefile.isOleFile(filepath):
        return {"error": "Not a valid OLE file", "text": "", "char_count": 0}
    ole = olefile.OleFileIO(filepath)
    text = ""
    if ole.exists("WordDocument"):
        # Try to extract raw text from the Word stream
        word_stream = ole.openstream("WordDocument").read()
        # Extract printable ASCII/UTF text (rough but works for most .doc)
        chars = []
        for b in word_stream:
            if 32 <= b < 127 or b in (10, 13, 9):
                chars.append(chr(b))
            elif len(chars) > 0 and chars[-1] != " ":
                chars.append(" ")
        text = "".join(chars).strip()
        # Clean up excessive whitespace
        import re
        text = re.sub(r" {3,}", "  ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
    ole.close()
    return {"text": text[:10000], "char_count": len(text)}


def parse_rtf(filepath: str) -> dict:
    """Extract text from RTF files."""
    from striprtf.striprtf import rtf_to_text
    with open(filepath, "r", errors="replace") as f:
        rtf_content = f.read()
    text = rtf_to_text(rtf_content)
    return {"text": text, "char_count": len(text)}


def parse_msg(filepath: str) -> dict:
    """Extract email data from Outlook .msg files."""
    import extract_msg
    msg = extract_msg.Message(filepath)
    result = {
        "from": msg.sender or "",
        "to": msg.to or "",
        "subject": msg.subject or "",
        "date": str(msg.date) if msg.date else "",
        "body": (msg.body or "")[:5000],
        "text": f"From: {msg.sender}\nTo: {msg.to}\nSubject: {msg.subject}\nDate: {msg.date}\n\n{(msg.body or '')[:5000]}",
        "char_count": len(msg.body or ""),
        "attachment_count": len(msg.attachments) if msg.attachments else 0,
        "attachments": [a.longFilename or a.shortFilename or "unnamed"
                       for a in (msg.attachments or [])],
    }
    msg.close()
    return result


def parse_xml(filepath: str) -> dict:
    """Extract from XML files."""
    import xml.etree.ElementTree as ET
    tree = ET.parse(filepath)
    root = tree.getroot()
    text = ET.tostring(root, encoding="unicode", method="text")[:5000]
    return {"root_tag": root.tag, "text": text, "char_count": len(text)}


def parse_text(filepath: str) -> dict:
    """Read plain text / CSV / JSON files."""
    with open(filepath, "r", errors="replace") as f:
        text = f.read(10000)
    return {"text": text, "char_count": len(text)}


def get_parser(ext: str):
    """Get the right parser for a file extension, handling legacy formats."""
    PARSERS = {
        ".xlsx": parse_xlsx,
        ".xls": parse_xls_legacy,  # Fixed! Uses xlrd
        ".docx": parse_docx,
        ".doc": parse_doc_legacy,  # Fixed! Uses olefile
        ".rtf": parse_rtf,
        ".msg": parse_msg, ".eml": parse_msg,
        ".xml": parse_xml,
        ".txt": parse_text, ".csv": parse_text, ".json": parse_text,
    }
    return PARSERS.get(ext)


# ─── PDF Classification (Phase 2: VECTOR) ─────────────────────────

def classify_and_extract_pdf(filepath: str) -> tuple[Lane, dict]:
    """Classify PDF as vector/scanned AND extract text if vector."""
    import fitz
    doc = fitz.open(filepath)
    total_chars = 0
    page_count = len(doc)
    full_text_parts = []

    for i, page in enumerate(doc):
        text = page.get_text()
        total_chars += len(text)
        full_text_parts.append(text)

    doc.close()

    avg_chars = total_chars / max(page_count, 1)
    full_text = "\n\n--- PAGE BREAK ---\n\n".join(full_text_parts)

    if avg_chars >= VECTOR_PDF_CHAR_THRESHOLD:
        lane = Lane.VECTOR
    else:
        lane = Lane.OCR

    metadata = {
        "page_count": page_count,
        "total_chars": total_chars,
        "avg_chars_per_page": round(avg_chars, 1),
    }
    return lane, metadata, full_text


# ─── Mistral OCR (Phase 3) ────────────────────────────────────────

def file_to_base64_data_url(filepath: str) -> tuple[str, str]:
    """Convert file to base64 data URL for Mistral API."""
    ext = Path(filepath).suffix.lower()
    mime_map = {
        ".pdf": "application/pdf",
        ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
        ".png": "image/png",
        ".bmp": "image/bmp",
        ".tiff": "image/tiff", ".tif": "image/tiff",
        ".doc": "application/msword",
        ".xls": "application/vnd.ms-excel",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    }
    mime = mime_map.get(ext, "application/octet-stream")
    with open(filepath, "rb") as f:
        data = base64.b64encode(f.read()).decode("ascii")
    return f"data:{mime};base64,{data}", mime


def ocr_with_mistral(filepath: str, api_key: str) -> dict:
    """
    Call Mistral OCR API for a single file.
    Returns: {"text": ..., "pages": ..., "char_count": ..., "model": "mistral-ocr-latest"}
    """
    data_url, mime = file_to_base64_data_url(filepath)

    # Determine document type for the API
    ext = Path(filepath).suffix.lower()
    image_exts = {".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif"}
    doc_type = "image_url" if ext in image_exts else "document_url"

    payload = {
        "model": "mistral-ocr-latest",
        "document": {
            "type": doc_type,
            doc_type: data_url,
        },
    }

    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }

    with httpx.Client(timeout=OCR_TIMEOUT) as client:
        resp = client.post(MISTRAL_OCR_ENDPOINT, json=payload, headers=headers)
        resp.raise_for_status()
        result = resp.json()

    # Extract text from Mistral OCR response
    pages = result.get("pages", [])
    all_text = []
    for page in pages:
        markdown = page.get("markdown", "")
        all_text.append(markdown)

    full_text = "\n\n--- PAGE BREAK ---\n\n".join(all_text)

    return {
        "text": full_text,
        "page_count": len(pages),
        "char_count": len(full_text),
        "model": "mistral-ocr-latest",
        "raw_response": result,
    }


# ─── Project Detection ─────────────────────────────────────────────

def detect_project_info(filepath: str, root_folder: str) -> tuple[str, str]:
    """Extract project folder name and subfolder from path."""
    rel = os.path.relpath(filepath, root_folder)
    parts = Path(rel).parts
    project = ""
    subfolder = ""
    for i, part in enumerate(parts):
        if part.startswith(("EH-", "EH ", "EHS", "SA ", "SA-", "OTH")):
            project = part
            if i + 1 < len(parts) - 1:
                subfolder = "/".join(parts[i + 1:-1])
            break
    if not project and len(parts) >= 3:
        project = parts[1] if len(parts) > 2 else parts[0]
        subfolder = "/".join(parts[2:-1]) if len(parts) > 3 else ""
    return project, subfolder


def extract_client_from_project(project_id: str) -> str:
    """Parse client name from project folder name like 'EH-02-26 ALBA FIT'."""
    parts = project_id.split(" ", 1)
    if len(parts) > 1:
        # Remove common suffixes like FIT, TIT, LIT, TW, AIT, DPIT, etc.
        name = parts[1]
        suffixes = ["FIT", "TIT", "LIT", "TW", "AIT", "DPIT", "PP", "BTU",
                    "CEMS", "SP", "EHS", "EH", "PH", "CL2", "TOC"]
        for suf in suffixes:
            if name.endswith(f" {suf}"):
                name = name[: -(len(suf) + 1)].strip()
        return name.strip()
    return project_id


# ─── Format Helpers ────────────────────────────────────────────────

def format_size(n: int) -> str:
    for u in ["B", "KB", "MB", "GB"]:
        if n < 1024:
            return f"{n:.1f} {u}"
        n /= 1024
    return f"{n:.1f} TB"


# ─── Main Pipeline ─────────────────────────────────────────────────

def run_pipeline(
    folder_path: str,
    mistral_key: str = "",
    do_ocr: bool = False,
    ocr_limit: int = 0,
) -> dict:
    """
    Full E2E pipeline: Discover → Classify → Parse → OCR → Merge.
    """
    all_files: list[FileResult] = []
    project_files: dict[str, list[FileResult]] = defaultdict(list)
    timings = {}

    # ── Phase 0: Discovery ──────────────────────────────────────
    t0 = time.time()
    print(f"\n{'='*70}")
    print(f"  ASYMMETRICA E2E OCR PIPELINE")
    print(f"  Folder: {folder_path}")
    print(f"  OCR:    {'ENABLED (Mistral)' if do_ocr else 'DISABLED (triage only)'}")
    print(f"{'='*70}")

    print(f"\n[Phase 0] Discovering files...")
    for dirpath, dirnames, filenames in os.walk(folder_path):
        for filename in filenames:
            filepath = os.path.join(dirpath, filename)
            try:
                size = os.path.getsize(filepath)
            except OSError:
                continue
            ext = Path(filename).suffix.lower()
            project, subfolder = detect_project_info(filepath, folder_path)
            result = FileResult(
                path=filepath, filename=filename, extension=ext,
                size_bytes=size, lane="", project_folder=project,
                subfolder=subfolder,
            )
            all_files.append(result)
            project_files[project].append(result)

    timings["discovery"] = time.time() - t0
    print(f"  Found {len(all_files)} files in {len(project_files)} projects ({timings['discovery']:.1f}s)")

    # ── Phase 1: Classification ─────────────────────────────────
    t1 = time.time()
    print(f"\n[Phase 1] Classifying files...")
    lane_counts = Counter()
    pdf_classify_count = 0

    for result in all_files:
        ext = result.extension
        if ext == ".pdf":
            try:
                lane, metadata, full_text = classify_and_extract_pdf(result.path)
                result.lane = lane.value
                result.page_count = metadata["page_count"]
                result.char_count = metadata["total_chars"]
                result.metadata = metadata
                if lane == Lane.VECTOR:
                    result.extracted_text = full_text
                    result.extracted_text_preview = full_text[:200]
                    result.confidence = 1.0
                pdf_classify_count += 1
            except Exception as e:
                result.lane = Lane.OCR.value
                result.error = str(e)
        elif ext in EXT_TO_LANE:
            mapped = EXT_TO_LANE[ext]
            result.lane = (mapped or Lane.SKIP).value
        else:
            result.lane = Lane.SKIP.value
        lane_counts[result.lane] += 1

    timings["classification"] = time.time() - t1
    print(f"  Classification done ({pdf_classify_count} PDFs analyzed) in {timings['classification']:.1f}s")
    for lane_val in [Lane.FREE.value, Lane.VECTOR.value, Lane.OCR.value, Lane.UNPACK.value, Lane.SKIP.value]:
        label = lane_val.replace("LANE_", "")
        print(f"    {label:>8}: {lane_counts.get(lane_val, 0):>4} files")

    # ── Phase 2: Free parsing ───────────────────────────────────
    t2 = time.time()
    free_files = [f for f in all_files if f.lane == Lane.FREE.value]
    print(f"\n[Phase 2] Parsing {len(free_files)} free-lane files...")
    parse_ok = 0
    parse_err = 0

    for i, result in enumerate(free_files):
        parser = get_parser(result.extension)
        if parser:
            try:
                parsed = parser(result.path)
                if "error" in parsed and not parsed.get("text"):
                    result.error = parsed["error"]
                    parse_err += 1
                else:
                    result.char_count = parsed.get("char_count", 0)
                    result.extracted_text = parsed.get("text", "")[:10000]
                    result.extracted_text_preview = result.extracted_text[:200]
                    result.confidence = 1.0
                    result.metadata = {k: v for k, v in parsed.items()
                                       if k not in ("text", "body", "raw_response")}
                    parse_ok += 1
            except Exception as e:
                result.error = str(e)
                parse_err += 1
        if (i + 1) % 100 == 0:
            print(f"    {i+1}/{len(free_files)} parsed...")

    timings["free_parse"] = time.time() - t2
    print(f"  Parsed {parse_ok} OK, {parse_err} errors in {timings['free_parse']:.1f}s")

    # ── Phase 2c: Fallback — promote failed free-parse to OCR ──
    promoted = 0
    for result in all_files:
        if result.lane == Lane.FREE.value and result.error and result.char_count == 0:
            result.lane = Lane.OCR.value
            result.metadata["fallback_reason"] = f"Free parse failed: {result.error}"
            result.error = ""  # Clear error — OCR will try
            lane_counts[Lane.FREE.value] -= 1
            lane_counts[Lane.OCR.value] += 1
            promoted += 1
    if promoted:
        print(f"  >> Promoted {promoted} failed files to OCR lane (fallback)")

    # ── Phase 3: Mistral OCR ────────────────────────────────────
    ocr_files = [f for f in all_files if f.lane == Lane.OCR.value]
    ocr_ok = 0
    ocr_err = 0
    ocr_cost = 0.0

    if do_ocr and mistral_key and ocr_files:
        t3 = time.time()
        files_to_ocr = ocr_files
        if ocr_limit > 0:
            files_to_ocr = ocr_files[:ocr_limit]
        print(f"\n[Phase 3] OCR via Mistral on {len(files_to_ocr)} files (of {len(ocr_files)} total)...")

        for i, result in enumerate(files_to_ocr):
            # Skip very large files
            if result.size_bytes > MAX_FILE_SIZE_FOR_OCR:
                result.error = f"Skipped: file too large ({format_size(result.size_bytes)} > 50MB)"
                ocr_err += 1
                continue

            try:
                t_ocr = time.time()
                ocr_result = ocr_with_mistral(result.path, mistral_key)
                result.extracted_text = ocr_result["text"][:10000]
                result.extracted_text_preview = ocr_result["text"][:200]
                result.char_count = ocr_result["char_count"]
                result.page_count = ocr_result["page_count"]
                result.confidence = 0.85  # Default OCR confidence
                result.ocr_model = ocr_result["model"]
                result.processing_time_ms = (time.time() - t_ocr) * 1000

                # Cost tracking (~$0.001/page estimate)
                page_cost = max(ocr_result["page_count"], 1) * 0.001
                ocr_cost += page_cost
                ocr_ok += 1

                print(f"    [{i+1}/{len(files_to_ocr)}] {result.filename[:50]:<50} "
                      f"{ocr_result['char_count']:>6} chars  "
                      f"{ocr_result['page_count']:>2} pages  "
                      f"${page_cost:.3f}  "
                      f"{result.processing_time_ms:.0f}ms")

            except httpx.HTTPStatusError as e:
                result.error = f"Mistral API {e.response.status_code}: {e.response.text[:200]}"
                ocr_err += 1
                print(f"    [{i+1}/{len(files_to_ocr)}] ERROR: {result.filename[:40]} — {result.error[:60]}")
            except Exception as e:
                result.error = f"OCR error: {str(e)[:200]}"
                ocr_err += 1
                print(f"    [{i+1}/{len(files_to_ocr)}] ERROR: {result.filename[:40]} — {result.error[:60]}")

            # Rate limit: small delay between calls
            time.sleep(MISTRAL_RATE_LIMIT_DELAY)

        timings["ocr"] = time.time() - t3
        print(f"  OCR done: {ocr_ok} OK, {ocr_err} errors, ${ocr_cost:.3f} total in {timings['ocr']:.1f}s")
    else:
        timings["ocr"] = 0
        if not do_ocr:
            print(f"\n[Phase 3] OCR SKIPPED (use --ocr to enable). {len(ocr_files)} files would be processed.")
        elif not mistral_key:
            print(f"\n[Phase 3] OCR SKIPPED — no Mistral API key provided.")

    # ── Phase 3b: Asymmetrica OCR Engines (post-processing) ───
    t3b = time.time()
    files_with_text = [f for f in all_files if f.extracted_text and len(f.extracted_text) > 10]
    print(f"\n[Phase 3b] Asymmetrica OCR Engines on {len(files_with_text)} extracted files...")

    from asymm_ocr_kernels import process_hybrid

    engine_corrections = 0
    engine_patterns = 0
    quality_before = []
    quality_after = []
    engine_details = []

    for i, result in enumerate(files_with_text):
        try:
            hybrid = process_hybrid(result.extracted_text)

            # Track quality change
            old_conf = result.confidence
            new_conf = hybrid.overall_quality
            quality_before.append(old_conf)
            quality_after.append(new_conf)

            # Update the file result with processed text
            result.extracted_text = hybrid.processed_text[:10000]
            result.extracted_text_preview = hybrid.processed_text[:200]
            result.confidence = new_conf

            # Store engine metadata
            result.metadata["asymm_engines"] = {
                "normalizer_quality": round(hybrid.normalizer_quality, 3),
                "vedic_confidence": round(hybrid.vedic_confidence, 3),
                "coherence_score": round(hybrid.coherence_score, 3),
                "overall_quality": round(hybrid.overall_quality, 3),
                "corrections": hybrid.corrections_count,
                "patterns": hybrid.patterns_found,
                "steps": hybrid.processing_steps,
            }

            engine_corrections += hybrid.corrections_count
            engine_patterns += hybrid.patterns_found

            # Track interesting findings
            if hybrid.corrections_count > 0 or hybrid.patterns_found > 0:
                engine_details.append({
                    "filename": result.filename,
                    "corrections": hybrid.corrections_count,
                    "patterns": hybrid.patterns_found,
                    "quality_delta": round(new_conf - old_conf, 3),
                    "pattern_details": hybrid.patterns[:3],  # Top 3
                    "correction_details": hybrid.corrections[:5],  # Top 5
                })

        except Exception as e:
            pass  # Don't fail the pipeline on engine errors

        if (i + 1) % 200 == 0:
            print(f"    {i+1}/{len(files_with_text)} processed...")

    timings["asymm_engines"] = time.time() - t3b

    avg_before = sum(quality_before) / max(len(quality_before), 1)
    avg_after = sum(quality_after) / max(len(quality_after), 1)
    files_improved = sum(1 for b, a in zip(quality_before, quality_after) if a > b)

    print(f"  Asymmetrica engines done in {timings['asymm_engines']:.1f}s")
    print(f"    Total corrections:  {engine_corrections}")
    print(f"    Total patterns:     {engine_patterns}")
    print(f"    Avg quality before: {avg_before:.3f}")
    print(f"    Avg quality after:  {avg_after:.3f}")
    print(f"    Files improved:     {files_improved}/{len(files_with_text)}")

    # ── Phase 4: Per-Project Merge ──────────────────────────────
    t4 = time.time()
    print(f"\n[Phase 4] Merging into project records...")
    projects: dict[str, ProjectRecord] = {}

    for proj_name, files in project_files.items():
        client = extract_client_from_project(proj_name)
        record = ProjectRecord(
            project_id=proj_name,
            client_name=client,
            file_count=len(files),
            total_size_bytes=sum(f.size_bytes for f in files),
            subfolders=list(set(f.subfolder for f in files if f.subfolder)),
        )

        confidences = []
        for f in files:
            file_entry = {
                "filename": f.filename,
                "extension": f.extension,
                "lane": f.lane,
                "subfolder": f.subfolder,
                "char_count": f.char_count,
                "confidence": f.confidence,
                "preview": f.extracted_text_preview[:150],
            }

            # Categorize into the right bucket
            fname_lower = f.filename.lower()
            if f.extension in (".msg", ".eml"):
                record.emails.append(file_entry)
            elif "costing" in fname_lower and f.extension in (".xlsx", ".xls"):
                record.costing_data.append(file_entry)
            elif "commercial" in fname_lower:
                record.commercial_offers.append(file_entry)
            elif "technical" in fname_lower:
                record.technical_specs.append(file_entry)
            elif f.subfolder and "rfq" in f.subfolder.lower():
                record.rfq_documents.append(file_entry)
            elif f.lane == Lane.OCR.value and f.extracted_text:
                record.ocr_results.append(file_entry)
            else:
                record.files.append(file_entry)

            if f.confidence > 0 and f.char_count > 0:
                confidences.append(f.confidence)

        record.confidence_avg = (
            sum(confidences) / len(confidences) if confidences else 0.0
        )
        projects[proj_name] = record

    timings["merge"] = time.time() - t4
    timings["total"] = time.time() - t0
    print(f"  Merged {len(projects)} project records in {timings['merge']:.1f}s")

    # ── Summary ─────────────────────────────────────────────────
    total_extracted = sum(1 for f in all_files if f.char_count > 0)
    total_chars = sum(f.char_count for f in all_files)
    errors = [f for f in all_files if f.error]

    print(f"\n{'='*70}")
    print(f"  PIPELINE COMPLETE")
    print(f"{'='*70}")
    print(f"  Total files:          {len(all_files)}")
    print(f"  Successfully extracted:{total_extracted:>5} ({total_extracted/max(len(all_files),1)*100:.1f}%)")
    print(f"  Total text extracted:  {total_chars:>10,} characters")
    print(f"  Errors:               {len(errors):>5}")
    print(f"  OCR cost:              ${ocr_cost:.3f}")
    print(f"  Total time:            {timings['total']:.1f}s")
    print(f"  Throughput:            {len(all_files)/max(timings['total'],0.001):.0f} files/sec")

    free_count = lane_counts.get(Lane.FREE.value, 0) + lane_counts.get(Lane.VECTOR.value, 0)
    print(f"\n  Asymmetrica savings:")
    print(f"    Files parsed for $0: {free_count}/{len(all_files)} ({free_count/max(len(all_files),1)*100:.1f}%)")
    print(f"    OCR cost:            ${ocr_cost:.3f}")
    print(f"    vs brute-force:      ${len(all_files) * 0.01:.2f}")
    if ocr_cost > 0:
        print(f"    Savings:             {(1 - ocr_cost/(len(all_files)*0.01))*100:.0f}%")
    print(f"{'='*70}\n")

    # Build export
    return {
        "folder": folder_path,
        "total_files": len(all_files),
        "total_size_bytes": sum(f.size_bytes for f in all_files),
        "total_projects": len(projects),
        "total_chars_extracted": total_chars,
        "lane_counts": dict(lane_counts),
        "ocr_cost_usd": round(ocr_cost, 4),
        "asymm_engines": {
            "total_corrections": engine_corrections,
            "total_patterns": engine_patterns,
            "avg_quality_before": round(avg_before, 4),
            "avg_quality_after": round(avg_after, 4),
            "files_improved": files_improved,
            "files_processed": len(files_with_text),
            "interesting_findings": engine_details[:30],
        },
        "timings": {k: round(v, 2) for k, v in timings.items()},
        "errors": [
            {"filename": f.filename, "path": f.path, "error": f.error}
            for f in errors
        ],
        "projects": {
            pid: {
                "project_id": p.project_id,
                "client_name": p.client_name,
                "file_count": p.file_count,
                "total_size_bytes": p.total_size_bytes,
                "subfolders": p.subfolders,
                "confidence_avg": round(p.confidence_avg, 3),
                "costing_sheets": len(p.costing_data),
                "commercial_offers": len(p.commercial_offers),
                "technical_specs": len(p.technical_specs),
                "emails": len(p.emails),
                "rfq_docs": len(p.rfq_documents),
                "ocr_results": len(p.ocr_results),
                "costing_data": p.costing_data,
                "commercial_offers_data": p.commercial_offers,
                "technical_specs_data": p.technical_specs,
                "email_data": p.emails,
                "rfq_data": p.rfq_documents,
                "ocr_data": p.ocr_results,
                "other_files": p.files,
            }
            for pid, p in sorted(projects.items())
        },
        "files": [
            {
                "path": f.path,
                "filename": f.filename,
                "extension": f.extension,
                "size_bytes": f.size_bytes,
                "lane": f.lane,
                "project": f.project_folder,
                "subfolder": f.subfolder,
                "char_count": f.char_count,
                "page_count": f.page_count,
                "confidence": f.confidence,
                "ocr_model": f.ocr_model,
                "preview": f.extracted_text_preview[:150],
                "full_text": f.extracted_text[:2000],  # Up to 2KB for engine processing
                "error": f.error,
                "time_ms": round(f.processing_time_ms, 1),
            }
            for f in all_files
        ],
    }


# ─── CLI ──────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Asymmetrica E2E OCR Pipeline"
    )
    parser.add_argument("--folder", required=True, help="Root folder to process")
    parser.add_argument("--mistral-key", default=os.environ.get("MISTRAL_API_KEY", ""),
                       help="Mistral API key")
    parser.add_argument("--ocr", action="store_true", help="Enable OCR (Phase 3)")
    parser.add_argument("--ocr-limit", type=int, default=0,
                       help="Max files to OCR (0=unlimited, useful for testing)")
    parser.add_argument("--json", default=None, help="Output JSON path")

    args = parser.parse_args()

    if not os.path.isdir(args.folder):
        print(f"ERROR: Folder not found: {args.folder}")
        sys.exit(1)

    result = run_pipeline(
        folder_path=args.folder,
        mistral_key=args.mistral_key,
        do_ocr=args.ocr,
        ocr_limit=args.ocr_limit,
    )

    json_path = args.json or os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        "pipeline_report.json"
    )
    with open(json_path, "w", encoding="utf-8") as fp:
        json.dump(result, fp, indent=2, ensure_ascii=False, default=str)
    print(f"Full report saved to: {json_path}")


if __name__ == "__main__":
    main()

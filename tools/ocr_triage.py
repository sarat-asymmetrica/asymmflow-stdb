"""
OCR Triage Pipeline — Asymmetrica Optimal Architecture
======================================================
Walks a document folder, classifies every file into processing lanes,
extracts data from free-parse files, and produces a full triage report.

Lanes:
  LANE_FREE     — Direct parse (XLSX, DOCX, RTF, MSG, XML, CSV, TXT)
  LANE_VECTOR   — Vector PDF (PyMuPDF text extraction, no OCR needed)
  LANE_OCR      — Scanned PDF / images (needs actual OCR)
  LANE_UNPACK   — Archives (ZIP) — extract and re-triage
  LANE_SKIP     — Video, HTML, unsupported (tag but don't process)

Philosophy: "The cheapest OCR call is the one you never make."

Usage:
  python ocr_triage.py --folder "path/to/offers" [--ocr] [--mistral-key KEY]
"""

import argparse
import hashlib
import json
import os
import sys
import time
from collections import Counter, defaultdict
from dataclasses import dataclass, field, asdict
from enum import Enum
from pathlib import Path
from typing import Optional


# ─── Lane Classification ───────────────────────────────────────────

class Lane(str, Enum):
    FREE = "LANE_FREE"         # Direct parse, $0
    VECTOR = "LANE_VECTOR"     # Vector PDF, $0
    OCR = "LANE_OCR"           # Needs actual OCR
    UNPACK = "LANE_UNPACK"     # Archive, unpack first
    SKIP = "LANE_SKIP"         # Unsupported / low value


# File extension → lane mapping
EXT_TO_LANE = {
    # FREE: Direct parseable formats
    ".xlsx": Lane.FREE, ".xls": Lane.FREE, ".csv": Lane.FREE,
    ".docx": Lane.FREE, ".doc": Lane.FREE,
    ".rtf": Lane.FREE,
    ".msg": Lane.FREE, ".eml": Lane.FREE,
    ".xml": Lane.FREE,
    ".txt": Lane.FREE, ".json": Lane.FREE,
    # OCR: Image formats (always need OCR)
    ".jpg": Lane.OCR, ".jpeg": Lane.OCR,
    ".png": Lane.OCR, ".bmp": Lane.OCR,
    ".tiff": Lane.OCR, ".tif": Lane.OCR,
    # UNPACK: Archives
    ".zip": Lane.UNPACK, ".rar": Lane.UNPACK, ".7z": Lane.UNPACK,
    # SKIP: Unsupported
    ".mp4": Lane.SKIP, ".mp3": Lane.SKIP, ".wav": Lane.SKIP,
    ".html": Lane.SKIP, ".htm": Lane.SKIP,
    ".exe": Lane.SKIP, ".dll": Lane.SKIP,
    ".ini": Lane.SKIP, ".log": Lane.SKIP,
    # PDF: Needs sub-classification (vector vs scanned)
    ".pdf": None,  # Special handling
}

# Minimum characters for a PDF page to be considered "vector" (has real text)
VECTOR_PDF_CHAR_THRESHOLD = 100


# ─── Data Structures ───────────────────────────────────────────────

@dataclass
class FileResult:
    path: str
    filename: str
    extension: str
    size_bytes: int
    lane: str
    project_folder: str = ""           # e.g., "EH-15-26 ALBA FIT"
    subfolder: str = ""                # e.g., "RFQ", "OFFER/REV-1"
    char_count: int = 0                # For vector PDFs / parsed text
    page_count: int = 0                # For PDFs
    confidence: float = 1.0            # 1.0 for direct parse, varies for OCR
    extracted_text_preview: str = ""   # First 200 chars
    metadata: dict = field(default_factory=dict)
    error: str = ""
    processing_time_ms: float = 0.0


@dataclass
class TriageReport:
    folder: str
    total_files: int = 0
    total_size_bytes: int = 0
    total_projects: int = 0
    lane_counts: dict = field(default_factory=dict)
    lane_sizes: dict = field(default_factory=dict)
    ext_counts: dict = field(default_factory=dict)
    project_summaries: dict = field(default_factory=dict)
    files: list = field(default_factory=list)
    errors: list = field(default_factory=list)
    timings: dict = field(default_factory=dict)
    estimated_ocr_cost_usd: float = 0.0
    estimated_ocr_time_sec: float = 0.0


# ─── File Parsers (Lane FREE) ──────────────────────────────────────

def parse_xlsx(filepath: str) -> dict:
    """Extract structure from Excel files — costing sheets, TQs, etc."""
    import openpyxl
    try:
        wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
        result = {
            "sheet_names": wb.sheetnames,
            "sheets": {},
        }
        total_chars = 0
        for sheet_name in wb.sheetnames[:5]:  # Cap at 5 sheets
            ws = wb[sheet_name]
            rows = []
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i >= 30:  # Preview first 30 rows
                    break
                row_data = [str(c) if c is not None else "" for c in row]
                rows.append(row_data)
                total_chars += sum(len(s) for s in row_data)
            result["sheets"][sheet_name] = {
                "rows_preview": rows,
                "max_row": ws.max_row,
                "max_col": ws.max_column,
            }
        wb.close()
        result["total_chars"] = total_chars
        return result
    except Exception as e:
        return {"error": str(e)}


def parse_docx(filepath: str) -> dict:
    """Extract text from Word documents."""
    import docx
    try:
        doc = docx.Document(filepath)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs)
        tables_data = []
        for table in doc.tables[:5]:  # Cap at 5 tables
            table_rows = []
            for row in table.rows[:20]:
                table_rows.append([cell.text.strip() for cell in row.cells])
            tables_data.append(table_rows)
        return {
            "text": text,
            "char_count": len(text),
            "paragraph_count": len(paragraphs),
            "table_count": len(doc.tables),
            "tables_preview": tables_data,
        }
    except Exception as e:
        return {"error": str(e)}


def parse_rtf(filepath: str) -> dict:
    """Extract text from RTF files."""
    from striprtf.striprtf import rtf_to_text
    try:
        with open(filepath, "r", errors="replace") as f:
            rtf_content = f.read()
        text = rtf_to_text(rtf_content)
        return {"text": text, "char_count": len(text)}
    except Exception as e:
        return {"error": str(e)}


def parse_msg(filepath: str) -> dict:
    """Extract email data from Outlook .msg files."""
    import extract_msg
    try:
        msg = extract_msg.Message(filepath)
        result = {
            "from": msg.sender or "",
            "to": msg.to or "",
            "subject": msg.subject or "",
            "date": str(msg.date) if msg.date else "",
            "body": (msg.body or "")[:2000],
            "char_count": len(msg.body or ""),
            "attachment_count": len(msg.attachments) if msg.attachments else 0,
            "attachments": [a.longFilename or a.shortFilename or "unnamed"
                           for a in (msg.attachments or [])],
        }
        msg.close()
        return result
    except Exception as e:
        return {"error": str(e)}


def parse_xml(filepath: str) -> dict:
    """Extract basic structure from XML files."""
    import xml.etree.ElementTree as ET
    try:
        tree = ET.parse(filepath)
        root = tree.getroot()
        text = ET.tostring(root, encoding="unicode", method="text")[:2000]
        return {
            "root_tag": root.tag,
            "child_count": len(list(root)),
            "text_preview": text,
            "char_count": len(text),
        }
    except Exception as e:
        return {"error": str(e)}


def parse_text(filepath: str) -> dict:
    """Read plain text files."""
    try:
        with open(filepath, "r", errors="replace") as f:
            text = f.read(10000)  # Cap at 10KB
        return {"text": text, "char_count": len(text)}
    except Exception as e:
        return {"error": str(e)}


# Extension → parser mapping
PARSERS = {
    ".xlsx": parse_xlsx, ".xls": parse_xlsx,
    ".docx": parse_docx, ".doc": parse_docx,
    ".rtf": parse_rtf,
    ".msg": parse_msg, ".eml": parse_msg,
    ".xml": parse_xml,
    ".txt": parse_text, ".csv": parse_text, ".json": parse_text,
}


# ─── PDF Sub-Classification ───────────────────────────────────────

def classify_pdf(filepath: str) -> tuple[Lane, dict]:
    """
    Classify a PDF as VECTOR (has extractable text) or OCR (scanned).
    Returns (lane, metadata).
    """
    import fitz  # PyMuPDF
    try:
        doc = fitz.open(filepath)
        total_chars = 0
        page_count = len(doc)
        page_texts = []

        for i, page in enumerate(doc):
            if i >= 10:  # Sample first 10 pages
                break
            text = page.get_text()
            total_chars += len(text)
            if i < 3:
                page_texts.append(text[:500])

        doc.close()

        # Decision: if average chars/page > threshold → vector PDF
        avg_chars_per_page = total_chars / max(page_count, 1)

        if avg_chars_per_page >= VECTOR_PDF_CHAR_THRESHOLD:
            lane = Lane.VECTOR
        else:
            lane = Lane.OCR

        metadata = {
            "page_count": page_count,
            "total_chars": total_chars,
            "avg_chars_per_page": round(avg_chars_per_page, 1),
            "text_preview": "\n---\n".join(page_texts)[:1000],
            "classification_reason": (
                f"Vector: {avg_chars_per_page:.0f} chars/page (threshold={VECTOR_PDF_CHAR_THRESHOLD})"
                if lane == Lane.VECTOR
                else f"Scanned: only {avg_chars_per_page:.0f} chars/page"
            ),
        }
        return lane, metadata

    except Exception as e:
        return Lane.OCR, {"error": str(e), "classification_reason": "Error reading PDF"}


# ─── Project Folder Detection ─────────────────────────────────────

def detect_project_info(filepath: str, root_folder: str) -> tuple[str, str]:
    """
    Extract project folder name and subfolder from file path.
    e.g., "Offers 2026/1-50/EH-15-26 ALBA FIT/OFFER/REV-1/quote.pdf"
    → project="EH-15-26 ALBA FIT", subfolder="OFFER/REV-1"
    """
    rel = os.path.relpath(filepath, root_folder)
    parts = Path(rel).parts

    project = ""
    subfolder = ""

    for i, part in enumerate(parts):
        # Project folders typically start with "EH-" or "EHS-" or "SA "
        if part.startswith(("EH-", "EH ", "EHS", "SA ")):
            project = part
            # Everything after project folder and before filename is subfolder
            if i + 1 < len(parts) - 1:
                subfolder = "/".join(parts[i + 1:-1])
            break

    # If no EH- pattern found, use first meaningful subfolder
    if not project and len(parts) >= 3:
        project = parts[1] if len(parts) > 2 else parts[0]
        subfolder = "/".join(parts[2:-1]) if len(parts) > 3 else ""

    return project, subfolder


# ─── Main Triage Engine ───────────────────────────────────────────

def triage_folder(folder_path: str, do_parse: bool = True) -> TriageReport:
    """
    Walk the folder, classify every file, optionally parse free-lane files.
    """
    report = TriageReport(folder=folder_path)
    lane_counts = Counter()
    lane_sizes = Counter()
    ext_counts = Counter()
    project_files = defaultdict(list)
    all_files = []

    start_time = time.time()

    # Phase 0: Discover all files
    print(f"\n🔍 Scanning: {folder_path}")
    print("─" * 60)

    for dirpath, dirnames, filenames in os.walk(folder_path):
        for filename in filenames:
            filepath = os.path.join(dirpath, filename)
            try:
                size = os.path.getsize(filepath)
            except OSError:
                continue

            ext = Path(filename).suffix.lower()
            project, subfolder = detect_project_info(filepath, folder_path)

            result = FileResult(
                path=filepath,
                filename=filename,
                extension=ext,
                size_bytes=size,
                lane="",
                project_folder=project,
                subfolder=subfolder,
            )
            all_files.append(result)
            ext_counts[ext] += 1

    discovery_time = time.time() - start_time
    print(f"📁 Found {len(all_files)} files in {discovery_time:.1f}s")

    # Phase 1: Classify each file
    print(f"\n⚡ Phase 1: Classifying files...")
    classify_start = time.time()

    for result in all_files:
        ext = result.extension

        if ext == ".pdf":
            # PDF needs sub-classification
            t0 = time.time()
            lane, metadata = classify_pdf(result.path)
            result.lane = lane.value
            result.page_count = metadata.get("page_count", 0)
            result.char_count = metadata.get("total_chars", 0)
            result.metadata = metadata
            result.extracted_text_preview = metadata.get("text_preview", "")[:200]
            result.processing_time_ms = (time.time() - t0) * 1000
        elif ext in EXT_TO_LANE:
            mapped = EXT_TO_LANE[ext]
            result.lane = (mapped or Lane.SKIP).value
        else:
            # Unknown extension
            result.lane = Lane.SKIP.value

        lane_counts[result.lane] += 1
        lane_sizes[result.lane] += result.size_bytes
        project_files[result.project_folder].append(result)

    classify_time = time.time() - classify_start
    print(f"✅ Classification done in {classify_time:.1f}s")

    # Phase 2: Parse free-lane files (if enabled)
    if do_parse:
        free_files = [f for f in all_files if f.lane == Lane.FREE.value]
        print(f"\n📖 Phase 2: Parsing {len(free_files)} free-lane files...")
        parse_start = time.time()
        parse_errors = 0

        for i, result in enumerate(free_files):
            parser = PARSERS.get(result.extension)
            if parser:
                t0 = time.time()
                try:
                    parsed = parser(result.path)
                    if "error" in parsed:
                        result.error = parsed["error"]
                        parse_errors += 1
                    else:
                        result.char_count = parsed.get("char_count", 0)
                        result.confidence = 1.0  # Direct parse = perfect
                        # Store preview
                        if "text" in parsed:
                            result.extracted_text_preview = parsed["text"][:200]
                        elif "sheets" in parsed:
                            # Excel: show first sheet first row
                            for sname, sdata in parsed["sheets"].items():
                                if sdata.get("rows_preview"):
                                    result.extracted_text_preview = str(sdata["rows_preview"][0])[:200]
                                    break
                        result.metadata = {
                            k: v for k, v in parsed.items()
                            if k not in ("text", "body")  # Don't store full text in metadata
                        }
                except Exception as e:
                    result.error = str(e)
                    parse_errors += 1
                result.processing_time_ms = (time.time() - t0) * 1000

            if (i + 1) % 50 == 0:
                print(f"   Parsed {i + 1}/{len(free_files)}...")

        parse_time = time.time() - parse_start
        print(f"✅ Parsing done in {parse_time:.1f}s ({parse_errors} errors)")
    else:
        parse_time = 0

    # Phase 2b: Extract text from vector PDFs
    vector_files = [f for f in all_files if f.lane == Lane.VECTOR.value]
    print(f"\n📄 Phase 2b: Vector PDF text already extracted during classification ({len(vector_files)} files)")

    # Compile report
    total_time = time.time() - start_time

    report.total_files = len(all_files)
    report.total_size_bytes = sum(f.size_bytes for f in all_files)
    report.total_projects = len(project_files)
    report.lane_counts = dict(lane_counts)
    report.lane_sizes = {k: v for k, v in lane_sizes.items()}
    report.ext_counts = dict(ext_counts.most_common())
    report.files = all_files
    report.errors = [f for f in all_files if f.error]
    report.timings = {
        "discovery_sec": round(discovery_time, 2),
        "classification_sec": round(classify_time, 2),
        "parsing_sec": round(parse_time, 2),
        "total_sec": round(total_time, 2),
    }

    # Estimate OCR costs (Mistral OCR pricing)
    ocr_files = [f for f in all_files if f.lane == Lane.OCR.value]
    ocr_pages = sum(max(f.page_count, 1) for f in ocr_files)
    report.estimated_ocr_cost_usd = ocr_pages * 0.001  # ~$0.001/page for Mistral
    report.estimated_ocr_time_sec = len(ocr_files) * 2.0  # ~2s per file estimate

    # Project summaries
    for proj_name, files in project_files.items():
        report.project_summaries[proj_name] = {
            "file_count": len(files),
            "total_size": sum(f.size_bytes for f in files),
            "lanes": dict(Counter(f.lane for f in files)),
            "extensions": dict(Counter(f.extension for f in files)),
            "has_costing_xlsx": any(
                "costing" in f.filename.lower() and f.extension in (".xlsx", ".xls")
                for f in files
            ),
            "has_commercial_offer": any(
                "commercial" in f.filename.lower() and f.extension in (".docx", ".pdf")
                for f in files
            ),
            "has_technical_offer": any(
                "technical" in f.filename.lower() and f.extension in (".docx", ".pdf")
                for f in files
            ),
            "subfolders": list(set(f.subfolder for f in files if f.subfolder)),
        }

    return report


# ─── Report Printer ───────────────────────────────────────────────

def format_size(bytes_val: int) -> str:
    """Human-readable file size."""
    for unit in ["B", "KB", "MB", "GB"]:
        if bytes_val < 1024:
            return f"{bytes_val:.1f} {unit}"
        bytes_val /= 1024
    return f"{bytes_val:.1f} TB"


def print_report(report: TriageReport):
    """Pretty-print the triage report."""

    print("\n" + "=" * 70)
    print("  ASYMMETRICA OCR TRIAGE REPORT")
    print("  \"The cheapest OCR call is the one you never make.\"")
    print("=" * 70)

    print(f"\n📁 Folder: {report.folder}")
    print(f"📊 Total files: {report.total_files}")
    print(f"💾 Total size: {format_size(report.total_size_bytes)}")
    print(f"🏗️  Projects detected: {report.total_projects}")

    # Lane breakdown
    print(f"\n{'─' * 60}")
    print("  LANE BREAKDOWN")
    print(f"{'─' * 60}")

    lane_labels = {
        Lane.FREE.value: ("🟢 FREE (direct parse)", "$0.00"),
        Lane.VECTOR.value: ("🔵 VECTOR PDF (PyMuPDF)", "$0.00"),
        Lane.OCR.value: ("🟠 OCR NEEDED (Mistral)", "~$0.001/pg"),
        Lane.UNPACK.value: ("📦 UNPACK (archive)", "—"),
        Lane.SKIP.value: ("⚪ SKIP (unsupported)", "—"),
    }

    total_free = 0
    for lane_val, (label, cost) in lane_labels.items():
        count = report.lane_counts.get(lane_val, 0)
        size = report.lane_sizes.get(lane_val, 0)
        pct = (count / max(report.total_files, 1)) * 100
        print(f"  {label}")
        print(f"    Files: {count:>6} ({pct:>5.1f}%)  |  Size: {format_size(size):>10}  |  Cost: {cost}")
        if lane_val in (Lane.FREE.value, Lane.VECTOR.value):
            total_free += count

    free_pct = (total_free / max(report.total_files, 1)) * 100
    print(f"\n  ✨ Files NOT needing OCR: {total_free}/{report.total_files} ({free_pct:.1f}%)")

    # OCR cost estimate
    ocr_count = report.lane_counts.get(Lane.OCR.value, 0)
    print(f"\n{'─' * 60}")
    print("  OCR COST ESTIMATE (Mistral OCR)")
    print(f"{'─' * 60}")
    print(f"  Files needing OCR:     {ocr_count}")
    print(f"  Estimated cost:        ${report.estimated_ocr_cost_usd:.2f}")
    print(f"  Estimated time:        {report.estimated_ocr_time_sec:.0f}s ({report.estimated_ocr_time_sec/60:.1f} min)")
    print(f"  vs. brute-force all:   ${report.total_files * 0.01:.2f} (10-20× more!)")

    # Extension breakdown
    print(f"\n{'─' * 60}")
    print("  FILE TYPES")
    print(f"{'─' * 60}")
    for ext, count in sorted(report.ext_counts.items(), key=lambda x: -x[1]):
        pct = (count / max(report.total_files, 1)) * 100
        bar = "█" * int(pct / 2)
        print(f"  {ext:>8}: {count:>5} ({pct:>5.1f}%) {bar}")

    # Top projects
    print(f"\n{'─' * 60}")
    print(f"  TOP PROJECTS (by file count)")
    print(f"{'─' * 60}")
    sorted_projects = sorted(
        report.project_summaries.items(),
        key=lambda x: -x[1]["file_count"]
    )
    for proj_name, summary in sorted_projects[:15]:
        fc = summary["file_count"]
        sz = format_size(summary["total_size"])
        has_cost = "💰" if summary["has_costing_xlsx"] else "  "
        has_comm = "📝" if summary["has_commercial_offer"] else "  "
        has_tech = "🔧" if summary["has_technical_offer"] else "  "
        print(f"  {has_cost}{has_comm}{has_tech} {proj_name:<40} {fc:>3} files  {sz:>10}")

    print(f"\n  Legend: 💰=Costing XLSX  📝=Commercial Offer  🔧=Technical Offer")

    # Errors
    if report.errors:
        print(f"\n{'─' * 60}")
        print(f"  ERRORS ({len(report.errors)} files)")
        print(f"{'─' * 60}")
        for f in report.errors[:10]:
            print(f"  ⚠️  {f.filename}: {f.error[:80]}")
        if len(report.errors) > 10:
            print(f"  ... and {len(report.errors) - 10} more")

    # Timing
    print(f"\n{'─' * 60}")
    print(f"  PERFORMANCE")
    print(f"{'─' * 60}")
    for label, val in report.timings.items():
        print(f"  {label:<25}: {val:.2f}s")
    files_per_sec = report.total_files / max(report.timings.get("total_sec", 1), 0.001)
    print(f"  {'throughput':<25}: {files_per_sec:.0f} files/sec")

    print(f"\n{'=' * 70}")
    print(f"  Done! Triage complete. 🎯")
    print(f"{'=' * 70}\n")


# ─── JSON Export ──────────────────────────────────────────────────

def export_report_json(report: TriageReport, output_path: str):
    """Export report to JSON for downstream processing."""

    export = {
        "folder": report.folder,
        "total_files": report.total_files,
        "total_size_bytes": report.total_size_bytes,
        "total_projects": report.total_projects,
        "lane_counts": report.lane_counts,
        "ext_counts": report.ext_counts,
        "estimated_ocr_cost_usd": report.estimated_ocr_cost_usd,
        "estimated_ocr_time_sec": report.estimated_ocr_time_sec,
        "timings": report.timings,
        "project_summaries": report.project_summaries,
        "files": [
            {
                "path": f.path,
                "filename": f.filename,
                "extension": f.extension,
                "size_bytes": f.size_bytes,
                "lane": f.lane,
                "project_folder": f.project_folder,
                "subfolder": f.subfolder,
                "char_count": f.char_count,
                "page_count": f.page_count,
                "confidence": f.confidence,
                "preview": f.extracted_text_preview[:100],
                "error": f.error,
                "processing_time_ms": round(f.processing_time_ms, 1),
            }
            for f in report.files
        ],
    }

    with open(output_path, "w", encoding="utf-8") as fp:
        json.dump(export, fp, indent=2, ensure_ascii=False, default=str)

    print(f"📄 JSON report saved to: {output_path}")


# ─── CLI ──────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Asymmetrica OCR Triage Pipeline — classify files before processing"
    )
    parser.add_argument(
        "--folder",
        required=True,
        help="Root folder to scan (e.g., 'Offers 2026/')"
    )
    parser.add_argument(
        "--no-parse",
        action="store_true",
        help="Skip parsing free-lane files (classification only)"
    )
    parser.add_argument(
        "--json",
        default=None,
        help="Export report to JSON file"
    )
    parser.add_argument(
        "--mistral-key",
        default=os.environ.get("MISTRAL_API_KEY", ""),
        help="Mistral API key for OCR (or set MISTRAL_API_KEY env var)"
    )

    args = parser.parse_args()

    if not os.path.isdir(args.folder):
        print(f"❌ Folder not found: {args.folder}")
        sys.exit(1)

    # Run triage
    report = triage_folder(args.folder, do_parse=not args.no_parse)

    # Print report
    print_report(report)

    # Export JSON if requested
    json_path = args.json
    if not json_path:
        # Default: save next to the script
        json_path = os.path.join(
            os.path.dirname(os.path.abspath(__file__)),
            "triage_report.json"
        )
    export_report_json(report, json_path)


if __name__ == "__main__":
    main()
